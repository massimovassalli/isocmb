from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os
import zipfile
import shutil

# Extract images from the Speakers.docx file
docx_path = 'ethicsinnovation/Speakers.docx'
output_dir = 'ethicsinnovation/speaker_images'

# Create output directory if it doesn't exist
os.makedirs(output_dir, exist_ok=True)

# Docx files are actually zip archives
# Extract the media folder
try:
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # List all files in the archive
        file_list = zip_ref.namelist()
        
        # Find all image files in word/media/
        image_files = [f for f in file_list if f.startswith('word/media/')]
        
        print(f"Found {len(image_files)} images in the document:")
        
        for img_file in image_files:
            # Extract the image
            img_data = zip_ref.read(img_file)
            
            # Get the filename
            img_name = os.path.basename(img_file)
            
            # Save the image
            output_path = os.path.join(output_dir, img_name)
            with open(output_path, 'wb') as f:
                f.write(img_data)
            
            print(f"  Extracted: {img_name}")
            
except Exception as e:
    print(f"Error extracting images: {e}")

# Also read the document structure
doc = Document(docx_path)
print(f"\n\nDocument structure:")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

# Check if speakers are in a table
if doc.tables:
    print("\nTable content:")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i+1}:")
        for row_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                print(f"    Row {row_idx+1}: {' | '.join(cells_text)}")
