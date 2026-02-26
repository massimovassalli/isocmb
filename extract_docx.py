from docx import Document
import os

# Read each docx file in the ethicsinnovation folder
files = ['Overview.docx', 'Partners.docx', 'Session Themes.docx', 'Speakers.docx']
base_path = 'ethicsinnovation'

for filename in files:
    filepath = os.path.join(base_path, filename)
    if os.path.exists(filepath):
        doc = Document(filepath)
        print(f"\n{'='*80}")
        print(f"FILE: {filename}")
        print('='*80)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        # Also check for tables
        for table in doc.tables:
            print("\n[TABLE FOUND]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                print(" | ".join(cells))
    else:
        print(f"\nFile not found: {filepath}")
